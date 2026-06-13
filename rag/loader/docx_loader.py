"""
文档加载器——DOCX 加载器

使用 python-docx 解析 Word 文档，提取段落文本和表格。
表格按行转为 Markdown 格式（| col1 | col2 |），保留结构信息。
"""

from pathlib import Path
from rag.datatypes import RawDocument
from .base import BaseLoader


class DocxLoader(BaseLoader):
    """DOCX 文档加载器"""

    def load(self, file_path: str) -> list[RawDocument]:
        path = Path(file_path)
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")

        doc = DocxDocument(str(path))
        # 提取段落文本
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 提取表格（转为 Markdown 格式）
        tables = self._extract_tables(doc)
        content = "\n".join(paragraphs)
        if tables:
            content += "\n\n" + "\n\n".join(tables)

        return [
            RawDocument(
                content=content,
                source=str(path),
                metadata={"file_name": path.name, "file_type": "docx"},
            )
        ]

    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def _extract_tables(self, doc) -> list[str]:
        """将 Word 表格按行转为 | col1 | col2 | 格式"""
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            tables.append("\n".join(rows))
        return tables
